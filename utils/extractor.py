"""
utils/extractor.py — Multi-format text extraction.

PDF   → pdfplumber (primary) → PyMuPDF (fallback) → OCR (last resort)
DOCX  → python-docx (paragraphs + tables + headers/footers)
Image → Advanced OpenCV pipeline → Tesseract (PSM 6 + PSM 3 + PSM 11)
        Special handling for dark-sidebar styled documents (resumes, etc.)
"""

import io
import logging

logger = logging.getLogger(__name__)

# ── Tesseract path for Windows ────────────────────────────────────────────────
try:
    import pytesseract
    import os
    # Common Windows Tesseract paths
    for path in [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Users\hp\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
    ]:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break
except Exception:
    pass


class TextExtractor:

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _pdf_pdfplumber(self, file_bytes: bytes) -> str:
        import pdfplumber
        texts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=3, y_tolerance=3)
                if t and t.strip():
                    texts.append(t.strip())
                else:
                    words = page.extract_words()
                    if words:
                        texts.append(" ".join(w["text"] for w in words))
        return "\n\n".join(texts)

    def _pdf_pymupdf(self, file_bytes: bytes) -> str:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = [p.get_text("text").strip() for p in doc if p.get_text("text").strip()]
        return "\n\n".join(parts)

    def _pdf_ocr(self, file_bytes: bytes) -> str:
        import fitz
        import pytesseract
        import numpy as np
        from PIL import Image
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        texts = []
        for page in doc:
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csGRAY)
            img = Image.fromarray(
                np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width)
            )
            t = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
            if t.strip():
                texts.append(t.strip())
        return "\n\n".join(texts)

    def _extract_pdf(self, file_bytes: bytes) -> str:
        text = ""
        try:
            text = self._pdf_pdfplumber(file_bytes)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        if len(text.strip()) < 100:
            try:
                text = self._pdf_pymupdf(file_bytes)
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        if len(text.strip()) < 50:
            try:
                text = self._pdf_ocr(file_bytes)
            except Exception as e:
                logger.warning(f"PDF OCR failed: {e}")
        return text

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, file_bytes: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                if para.style.name.startswith("Heading"):
                    parts.append(f"[HEADING] {t}")
                else:
                    parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        try:
            for section in doc.sections:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        parts.append(f"[HEADER] {para.text.strip()}")
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        parts.append(f"[FOOTER] {para.text.strip()}")
        except Exception:
            pass

        return "\n".join(parts)

    # ── IMAGE ─────────────────────────────────────────────────────────────────

    def _preprocess_standard(self, img_bgr):
        """Standard preprocessing for plain documents."""
        import cv2
        h, w = img_bgr.shape[:2]
        if max(h, w) < 1500:
            scale = 1500 / max(h, w)
            img_bgr = cv2.resize(img_bgr, None, fx=scale, fy=scale,
                                 interpolation=cv2.INTER_CUBIC)
        gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray = clahe.apply(gray)
        gray = cv2.GaussianBlur(gray, (3, 3), 0)
        thresh = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 31, 8
        )
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
        return cv2.morphologyEx(thresh, cv2.MORPH_CLOSE, kernel)

    def _preprocess_dark_sidebar(self, img_bgr):
        """
        Special preprocessing for styled documents with dark sidebars
        (e.g. resumes with dark left panel).
        Crops only the right/light portion for OCR.
        """
        import cv2
        import numpy as np

        gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
        h, w = gray.shape

        # Find where image becomes predominantly light (right portion)
        col_means = np.mean(gray, axis=0)
        light_cols = col_means > 180
        # Find first column where >60% of cols to the right are light
        split = 0
        for i in range(w // 4, w):
            if np.mean(light_cols[i:]) > 0.6:
                split = i
                break

        if split > w // 5:
            # Crop to light portion
            cropped = img_bgr[:, split:]
        else:
            cropped = img_bgr

        return self._preprocess_standard(cropped)

    def _deskew(self, thresh_img):
        """Fix slight rotation in scanned documents."""
        import cv2
        import numpy as np
        coords = np.column_stack(np.where(thresh_img > 0))
        if len(coords) < 10:
            return thresh_img
        angle = cv2.minAreaRect(coords)[-1]
        if angle < -45:
            angle = 90 + angle
        if abs(angle) < 0.5:
            return thresh_img
        (h, w) = thresh_img.shape
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        return cv2.warpAffine(thresh_img, M, (w, h),
                              flags=cv2.INTER_CUBIC,
                              borderMode=cv2.BORDER_REPLICATE)

    def _load_image(self, file_bytes: bytes):
        """Load image bytes → BGR numpy array."""
        import cv2
        import numpy as np
        from PIL import Image

        nparr = np.frombuffer(file_bytes, dtype=np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        if img is None:
            pil = Image.open(io.BytesIO(file_bytes)).convert("RGB")
            img = cv2.cvtColor(np.array(pil), cv2.COLOR_RGB2BGR)
        return img

    def _ocr_image(self, thresh_img) -> str:
        """Run Tesseract with multiple PSM modes, pick best result."""
        import pytesseract
        from PIL import Image

        pil = Image.fromarray(thresh_img)
        results = []
        for psm in (6, 3, 11):
            try:
                t = pytesseract.image_to_string(
                    pil, config=f"--oem 3 --psm {psm}"
                ).strip()
                if t:
                    results.append(t)
            except Exception:
                continue
        if not results:
            return ""
        return max(results, key=len)

    def _extract_image(self, file_bytes: bytes) -> str:
        img = self._load_image(file_bytes)

        # Strategy 1: standard preprocessing
        thresh1 = self._preprocess_standard(img)
        thresh1 = self._deskew(thresh1)
        text1 = self._ocr_image(thresh1)

        # Strategy 2: dark sidebar handling
        thresh2 = self._preprocess_dark_sidebar(img)
        thresh2 = self._deskew(thresh2)
        text2 = self._ocr_image(thresh2)

        # Pick the result with more text content
        return text1 if len(text1) >= len(text2) else text2

    # ── Router ────────────────────────────────────────────────────────────────

    def extract(self, file_bytes: bytes, file_type: str) -> str:
        ft = file_type.lower().strip()
        if ft == "pdf":
            return self._extract_pdf(file_bytes)
        elif ft == "docx":
            return self._extract_docx(file_bytes)
        elif ft == "image":
            return self._extract_image(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")